from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt


SOURCE_MODE_REAL = "retrieved_real_precedent"
SOURCE_MODE_LEGACY = "legacy_imported_precedent"
SOURCE_MODE_INTERNAL = "internal_fallback_profile"


def build_legal_query_docx(
    response_payload: dict[str, Any],
    request_context: dict[str, Any] | None = None,
) -> bytes:
    response = _as_dict(response_payload)
    request = _as_dict(request_context)
    reasoning = _as_dict(response.get("reasoning"))
    procedural_strategy = _as_dict(response.get("procedural_strategy"))
    case_struct = _as_dict(response.get("case_structure"))
    norm_reasoning = _as_dict(response.get("normative_reasoning"))
    question_result = _as_dict(response.get("question_engine_result"))
    case_theory = _as_dict(response.get("case_theory"))
    case_evaluation = _as_dict(response.get("case_evaluation"))
    conflict_evidence = _as_dict(response.get("conflict_evidence"))
    evidence_links = _as_dict(response.get("evidence_reasoning_links"))
    jurisprudence_analysis = _as_dict(response.get("jurisprudence_analysis"))
    case_profile = _resolve_case_profile(response)
    case_strategy = _resolve_case_strategy(response)
    legal_strategy = _as_dict(response.get("legal_strategy"))

    document = Document()
    _set_base_styles(document)

    query = _as_text(response.get("query")) or _as_text(request.get("query")) or "Consulta juridica"
    jurisdiction = _as_text(response.get("jurisdiction")) or _as_text(request.get("jurisdiction"))
    forum = _as_text(response.get("forum")) or _as_text(request.get("forum"))
    case_domain = _as_text(response.get("case_domain")) or _as_text(legal_strategy.get("case_domain")) or _as_text(case_profile.get("case_domain"))
    case_domains = _as_list(response.get("case_domains")) or _as_list(legal_strategy.get("case_domains")) or _as_list(case_profile.get("case_domains"))
    document_mode = _as_text(request.get("document_mode"))
    confidence = response.get("confidence")

    document.add_heading("AILEX - Exportacion de resultado juridico", level=0)
    document.add_paragraph(query)

    _add_kv_block(
        document,
        [
            ("Jurisdiccion", jurisdiction),
            ("Foro", forum),
            ("Dominio principal", case_domain),
            ("Dominios detectados", ", ".join(_format_any(item) for item in case_domains if _format_any(item))),
            ("Modo documental", document_mode),
            ("Confianza", _format_confidence(confidence)),
        ],
    )

    short_answer = (
        _as_text(case_strategy.get("strategic_narrative"))
        or _as_text(reasoning.get("short_answer"))
        or _as_text(reasoning.get("case_analysis"))
        or _as_text(reasoning.get("applied_analysis"))
    )
    if short_answer:
        _add_section(document, "Respuesta breve", [short_answer])

    legal_strategy_lines = _format_legal_strategy(
        case_strategy,
        case_domain=case_domain,
        case_domains=case_domains,
    )
    if legal_strategy_lines:
        _add_section(document, "Estrategia juridica estructurada", legal_strategy_lines)

    questions = _format_questions(question_result, case_strategy)
    if questions:
        _add_section(document, "Preguntas clave para completar el caso", questions)

    strategy_lines = _format_strategy(procedural_strategy)
    if strategy_lines:
        _add_section(document, "Estrategia procesal", strategy_lines)

    normative_lines = _format_visible_normative_sections(response, reasoning, norm_reasoning)
    for title, lines in normative_lines:
        _add_section(document, title, lines)

    case_facts = _as_list(case_struct.get("facts"))
    if case_facts:
        _add_section(document, "Hechos detectados", case_facts)

    main_claim = _as_text(case_struct.get("main_claim"))
    if main_claim:
        _add_section(document, "Pretension principal", [main_claim])

    case_missing = _as_list(case_struct.get("missing_information"))
    if case_missing:
        _add_section(document, "Informacion faltante", case_missing)

    case_risks = _as_list(case_struct.get("risks"))
    if case_risks:
        _add_section(document, "Riesgos del caso", case_risks)

    theory_lines = _format_case_theory(case_theory)
    if theory_lines:
        _add_section(document, "Teoria del caso", theory_lines)

    evaluation_lines = _format_case_evaluation(case_evaluation)
    if evaluation_lines:
        _add_section(document, "Evaluacion estrategica del caso", evaluation_lines)

    conflict_lines = _format_conflict_evidence(conflict_evidence)
    if conflict_lines:
        _add_section(document, "Conflicto y prueba", conflict_lines)

    evidence_link_lines = _format_evidence_reasoning_links(evidence_links)
    if evidence_link_lines:
        _add_section(document, "Trazabilidad probatoria", evidence_link_lines)

    jurisprudence_lines = _format_jurisprudence_analysis(jurisprudence_analysis)
    if jurisprudence_lines:
        _add_section(document, "Jurisprudencia relevante", jurisprudence_lines)

    warnings = _collect_warnings(response)
    if warnings:
        _add_section(document, "Advertencias", warnings)

    citations = _format_citations(_as_list(reasoning.get("citations_used")))
    if citations:
        _add_section(document, "Citas utilizadas", citations)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _set_base_styles(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)


def _add_kv_block(document: Document, rows: list[tuple[str, str]]) -> None:
    visible_rows = [(label, value) for label, value in rows if value]
    if not visible_rows:
        return

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in visible_rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _add_section(document: Document, title: str, items: list[str]) -> None:
    visible_items = [item for item in items if item]
    if not visible_items:
        return
    document.add_heading(title, level=1)
    for item in visible_items:
        document.add_paragraph(item, style="List Bullet")


def _format_strategy(strategy: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    next_steps = _as_list(strategy.get("next_steps"))
    if not next_steps and isinstance(strategy.get("steps"), list):
        next_steps = [_format_strategy_step(item) for item in strategy.get("steps") or []]
    risks = _as_list(strategy.get("risks"))
    missing = _as_list(strategy.get("missing_information") or strategy.get("missing_info"))

    for item in next_steps:
        text = _format_any(item)
        if text:
            lines.append(f"Paso procesal complementario: {text}")
    for item in risks:
        text = _format_any(item)
        if text:
            lines.append(f"Riesgo procesal complementario: {text}")
    for item in missing:
        text = _format_any(item)
        if text:
            lines.append(f"Informacion faltante procesal: {text}")
    return lines


def _format_legal_strategy(
    strategy: dict[str, Any],
    *,
    case_domain: str = "",
    case_domains: list[Any] | None = None,
) -> list[str]:
    lines: list[str] = []
    if not strategy:
        return lines

    visible_domains = [text for text in (_format_any(item) for item in (case_domains or [])) if text]
    if case_domain:
        lines.append(f"Dominio principal: {case_domain}")
    if visible_domains:
        lines.append(f"Dominios detectados: {', '.join(visible_domains)}")

    narrative = _as_text(strategy.get("strategic_narrative"))
    if narrative:
        lines.append(f"Estrategia: {narrative}")
    for item in _as_list(strategy.get("conflict_summary")):
        text = _format_any(item)
        if text:
            lines.append(f"Conflicto principal: {text}")
    for item in _as_list(strategy.get("recommended_actions")):
        text = _format_any(item)
        if text:
            lines.append(f"Accion recomendada: {text}")
    for item in _as_list(strategy.get("risk_analysis")):
        text = _format_any(item)
        if text:
            lines.append(f"Riesgo: {text}")
    for item in _as_list(strategy.get("procedural_focus")):
        text = _format_any(item)
        if text:
            lines.append(f"Foco procesal: {text}")
    for item in _as_list(strategy.get("secondary_domain_notes")):
        text = _format_any(item)
        if text:
            lines.append(f"Nota por dominio secundario: {text}")
    return lines


def _format_questions(question_result: dict[str, Any], case_strategy: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    for item in _as_list(case_strategy.get("critical_questions")):
        text = _format_any(item)
        if text:
            lines.append(text)
    for item in _as_list(question_result.get("critical_questions")):
        text = _format_any(item)
        if text and text not in lines:
            lines.append(text)
    for item in _as_list(question_result.get("questions")):
        text = _format_question_item(item)
        if text and text not in lines:
            lines.append(text)
    return lines


def _format_visible_normative_sections(
    response: dict[str, Any],
    reasoning: dict[str, Any],
    norm_reasoning: dict[str, Any],
) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []
    case_domain = _as_text(response.get("case_domain"))
    warnings_text = " ".join(_as_list(norm_reasoning.get("warnings"))).lower()
    summary_text = _as_text(norm_reasoning.get("summary")).lower()
    generic_normative = "fallback generico" in warnings_text or "razonamiento normativo generico" in summary_text

    normative_foundations = _format_normative_foundations(
        _as_list(reasoning.get("normative_foundations") or reasoning.get("normative_grounds"))
    )
    if normative_foundations and not (generic_normative and case_domain == "conflicto_patrimonial"):
        sections.append(("Fundamentos normativos", normative_foundations))

    applied_rules = _format_applied_rules(_as_list(norm_reasoning.get("applied_rules")))
    if applied_rules:
        title = "Reglas aplicadas"
        if generic_normative:
            title = "Normativa secundaria o de apoyo"
        sections.append((title, applied_rules))

    requirements = [_format_any(item) for item in _as_list(norm_reasoning.get("requirements")) if _format_any(item)]
    if requirements:
        title = "Requisitos legales"
        if generic_normative:
            title = "Requisitos a verificar"
        sections.append((title, requirements))

    unresolved = [_format_any(item) for item in _as_list(norm_reasoning.get("unresolved_issues")) if _format_any(item)]
    if unresolved:
        sections.append(("Cuestiones pendientes", unresolved))

    return sections


def _format_normative_foundations(items: list[Any]) -> list[str]:
    formatted: list[str] = []
    for item in items:
        if isinstance(item, str):
            formatted.append(item)
            continue
        if not isinstance(item, dict):
            continue
        label = _as_text(item.get("label")) or _as_text(item.get("title")) or _as_text(item.get("titulo"))
        source = _as_text(item.get("source_id")) or _as_text(item.get("source")) or _as_text(item.get("norma"))
        article = _as_text(item.get("article")) or _as_text(item.get("articulo"))
        excerpt = _as_text(item.get("texto")) or _as_text(item.get("text")) or _as_text(item.get("summary"))

        parts = []
        if label:
            parts.append(label)
        if source:
            parts.append(f"Fuente: {source}")
        if article:
            parts.append(f"Articulo: {article}")
        if excerpt:
            parts.append(excerpt)
        if parts:
            formatted.append(" | ".join(parts))
    return formatted


def _format_applied_rules(items: list[Any]) -> list[str]:
    formatted: list[str] = []
    for item in items:
        if isinstance(item, str):
            formatted.append(item)
            continue
        if not isinstance(item, dict):
            continue
        source = _as_text(item.get("source"))
        article = _as_text(item.get("article"))
        relevance = _as_text(item.get("relevance"))
        effect = _as_text(item.get("effect"))

        parts = []
        if source and article:
            parts.append(f"Art. {article} - {source}")
        elif article:
            parts.append(f"Art. {article}")
        if relevance:
            parts.append(relevance)
        if effect:
            parts.append(f"Efecto: {effect}")
        if parts:
            formatted.append(" | ".join(parts))
    return formatted


def _format_case_theory(case_theory: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    theory_primary = _as_text(case_theory.get("primary_theory"))
    theory_objective = _as_text(case_theory.get("objective"))
    if theory_primary:
        lines.append(f"Teoria principal: {theory_primary}")
    if theory_objective:
        lines.append(f"Objetivo: {theory_objective}")
    for item in _as_list(case_theory.get("alternative_theories")):
        text = _format_any(item)
        if text:
            lines.append(f"Teoria alternativa: {text}")
    for item in _as_list(case_theory.get("key_facts_supporting")):
        text = _format_any(item)
        if text:
            lines.append(f"Hecho clave: {text}")
    for item in _as_list(case_theory.get("likely_points_of_conflict")):
        text = _format_any(item)
        if text:
            lines.append(f"Punto de conflicto: {text}")
    for item in _as_list(case_theory.get("evidentiary_needs")):
        text = _format_any(item)
        if text:
            lines.append(f"Necesidad probatoria: {text}")
    for item in _as_list(case_theory.get("recommended_line_of_action")):
        text = _format_any(item)
        if text:
            lines.append(f"Linea recomendada: {text}")
    return lines


def _format_case_evaluation(case_evaluation: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    for label, key in (
        ("Fortaleza del caso", "case_strength"),
        ("Nivel de riesgo", "legal_risk_level"),
        ("Nivel de incertidumbre", "uncertainty_level"),
    ):
        value = _as_text(case_evaluation.get(key))
        if value:
            lines.append(f"{label}: {value}")
    for item in _as_list(case_evaluation.get("strategic_observations")):
        text = _format_any(item)
        if text:
            lines.append(f"Observacion estrategica: {text}")
    for item in _as_list(case_evaluation.get("possible_scenarios")):
        text = _format_any(item)
        if text:
            lines.append(f"Escenario posible: {text}")
    return lines


def _format_conflict_evidence(conflict_evidence: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    for label, key in (
        ("Nucleo del conflicto", "core_dispute"),
        ("Punto mas fuerte", "strongest_point"),
        ("Punto mas vulnerable", "most_vulnerable_point"),
    ):
        value = _as_text(conflict_evidence.get(key))
        if value:
            lines.append(f"{label}: {value}")
    for item in _as_list(conflict_evidence.get("critical_evidence_available")):
        text = _format_any(item)
        if text:
            lines.append(f"Prueba critica disponible: {text}")
    for item in _as_list(conflict_evidence.get("key_evidence_missing")):
        text = _format_any(item)
        if text:
            lines.append(f"Prueba critica faltante: {text}")
    for item in _as_list(conflict_evidence.get("probable_counterarguments")):
        text = _format_any(item)
        if text:
            lines.append(f"Contraargumento probable: {text}")
    for item in _as_list(conflict_evidence.get("recommended_evidence_actions")):
        text = _format_any(item)
        if text:
            lines.append(f"Accion recomendada sobre prueba: {text}")
    return lines


def _format_evidence_reasoning_links(evidence_links: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    summary = _as_text(evidence_links.get("summary"))
    if summary:
        lines.append(summary)
    level_labels = {"alto": "ALTO", "medio": "MEDIO", "bajo": "BAJO"}

    for link in _as_list(evidence_links.get("requirement_links")):
        if not isinstance(link, dict):
            continue
        requirement = _as_text(link.get("requirement"))
        if not requirement:
            continue
        support = _as_text(link.get("support_level"))
        source = _as_text(link.get("source"))
        article = _as_text(link.get("article"))
        note = _as_text(link.get("strategic_note"))
        support_label = level_labels.get(support, "BAJO")
        header = f"[{support_label}] {requirement}"
        if source and article:
            header = f"{header} - Art. {article} ({source})"
        lines.append(header)
        for item in _as_list(link.get("supporting_facts"))[:2]:
            text = _format_any(item)
            if text:
                lines.append(f"Hecho de apoyo: {text}")
        for item in _as_list(link.get("evidence_available"))[:2]:
            text = _format_any(item)
            if text:
                lines.append(f"Evidencia disponible: {text}")
        for item in _as_list(link.get("evidence_missing"))[:2]:
            text = _format_any(item)
            if text:
                lines.append(f"Evidencia faltante: {text}")
        if note:
            lines.append(f"Nota: {note}")
    for item in _as_list(evidence_links.get("critical_evidentiary_gaps")):
        text = _format_any(item)
        if text:
            lines.append(f"Brecha probatoria: {text}")
    for item in _as_list(evidence_links.get("strategic_warnings")):
        text = _format_any(item)
        if text:
            lines.append(f"Advertencia estrategica: {text}")
    return lines


def _format_jurisprudence_analysis(jurisprudence_analysis: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    source_summary = _as_text(jurisprudence_analysis.get("source_mode_summary"))
    source_quality = _as_text(jurisprudence_analysis.get("source_quality"))
    strength = _as_text(jurisprudence_analysis.get("jurisprudence_strength"))
    if source_quality:
        lines.append(f"Calidad de fuente: {_format_source_quality(source_quality)}")
    if strength:
        lines.append(f"Fuerza orientativa: {strength}")
    if source_summary:
        lines.append(f"Resumen de fuente: {source_summary}")

    for item in _as_list(jurisprudence_analysis.get("jurisprudence_highlights"))[:3]:
        if not isinstance(item, dict):
            continue
        block = _format_jurisprudence_highlight(item)
        if block:
            lines.append(block)
    return lines


def _format_jurisprudence_highlight(item: dict[str, Any]) -> str:
    parts: list[str] = []
    source_mode = _as_text(item.get("source_mode"))
    label = _label_for_source_mode(source_mode)
    case_name = _as_text(item.get("case_name"))
    court = _as_text(item.get("court"))
    year = _as_text(item.get("year"))
    criterion = _as_text(item.get("criterion"))
    strategic_use = _as_text(item.get("strategic_use"))
    if label:
        parts.append(label)
    if case_name:
        parts.append(f"Caso: {case_name}")
    if court and source_mode in {SOURCE_MODE_REAL, SOURCE_MODE_LEGACY}:
        parts.append(f"Tribunal: {court}")
    if year and year != "0":
        parts.append(f"Ano: {year}")
    if criterion:
        parts.append(f"Criterio: {criterion}")
    if strategic_use:
        parts.append(f"Utilidad estrategica: {strategic_use}")
    if source_mode == SOURCE_MODE_INTERNAL:
        parts.append("No constituye precedente real verificable del corpus.")
    return " | ".join(parts)


def _format_source_quality(value: str) -> str:
    mapping = {
        "real": "precedentes reales del corpus",
        "legacy": "precedentes legacy importados",
        "fallback": "fallback interno orientativo",
        "none": "sin base jurisprudencial suficiente",
    }
    return mapping.get(value, value)


def _label_for_source_mode(source_mode: str) -> str:
    if source_mode == SOURCE_MODE_REAL:
        return "Precedente real recuperado"
    if source_mode == SOURCE_MODE_LEGACY:
        return "Precedente legacy importado"
    if source_mode == SOURCE_MODE_INTERNAL:
        return "Perfil interno orientativo"
    return "Fuente jurisprudencial no especificada"


def _format_question_item(item: Any) -> str:
    if isinstance(item, str):
        return item.strip()
    if not isinstance(item, dict):
        return ""
    question = _as_text(item.get("question"))
    purpose = _as_text(item.get("purpose"))
    priority = _as_text(item.get("priority"))
    category = _as_text(item.get("category"))
    parts = []
    if priority:
        parts.append(f"[{priority.upper()}]")
    if category:
        parts.append(category)
    if question:
        parts.append(question)
    if purpose:
        parts.append(f"Objetivo: {purpose}")
    return " | ".join(parts)


def _format_citations(items: list[Any]) -> list[str]:
    return [text for text in (_format_any(item) for item in items) if text]


def _format_strategy_step(item: Any) -> str:
    if isinstance(item, str):
        return item.strip()
    if isinstance(item, dict):
        action = _as_text(item.get("action") or item.get("label") or item.get("title"))
        deadline = _as_text(item.get("deadline_hint"))
        if action and deadline:
            return f"{action} ({deadline})"
        return action
    return _format_any(item)


def _collect_warnings(response: dict[str, Any]) -> list[str]:
    seen: set[str] = set()
    items: list[str] = []
    for payload in (
        response,
        _as_dict(response.get("reasoning")),
        _as_dict(response.get("citation_validation")),
        _as_dict(response.get("hallucination_guard")),
        _as_dict(response.get("conflict_evidence")),
        _as_dict(response.get("evidence_reasoning_links")),
        _as_dict(response.get("jurisprudence_analysis")),
        _resolve_case_profile(response),
        _resolve_case_strategy(response),
    ):
        for raw in (
            *_as_list(payload.get("warnings")),
            *_as_list(payload.get("strategic_warnings")),
            payload.get("source_mode_summary"),
        ):
            text = _format_any(raw)
            if not text or text in seen:
                continue
            seen.add(text)
            items.append(text)
    return items


def _resolve_case_profile(response: dict[str, Any]) -> dict[str, Any]:
    direct = _as_dict(response.get("case_profile"))
    if direct:
        return direct
    return _as_dict(_as_dict(response.get("legal_strategy")).get("case_profile"))


def _resolve_case_strategy(response: dict[str, Any]) -> dict[str, Any]:
    direct = _as_dict(response.get("case_strategy"))
    if direct:
        return direct
    return _as_dict(_as_dict(response.get("legal_strategy")).get("case_strategy"))


def _format_confidence(value: Any) -> str:
    if isinstance(value, (int, float)):
        return f"{round(max(0.0, min(1.0, float(value))) * 100)}%"
    return ""


def _format_any(value: Any) -> str:
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        for key in ("label", "title", "titulo", "name", "description", "text", "article", "source_id"):
            text = _as_text(value.get(key))
            if text:
                return text
    return _as_text(value)


def _as_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int, float)):
        return str(value)
    return ""


def _as_list(value: Any) -> list[Any]:
    return value if isinstance(value, list) else []


def _as_dict(value: Any) -> dict[str, Any]:
    return value if isinstance(value, dict) else {}
